from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import math
import os
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from inference.plugins.qwen_endpoint import dashscope_base_url

logger = logging.getLogger(__name__)

DEFAULT_CHUNK_SIZE = 900
DEFAULT_CHUNK_OVERLAP = 120
DEFAULT_TOP_K = 5
DEFAULT_MAX_CONTEXT_CHARS = 4500
DEFAULT_MIN_SCORE = 0.25


@dataclass
class RAGIndexRequest:
    character_id: str
    character_dir: str
    source_id: str
    source_type: str
    title: str
    filename: str
    mime_type: str
    source_path: str


@dataclass
class RAGSearchRequest:
    character_id: str
    character_dir: str
    query: str
    top_k: int = DEFAULT_TOP_K
    max_context_chars: int = DEFAULT_MAX_CONTEXT_CHARS
    min_score: float = DEFAULT_MIN_SCORE


@dataclass
class RAGSearchResult:
    source_id: str
    source_type: str
    title: str
    filename: str
    content: str
    score: float


def _safe_collection_name(character_id: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9_-]+", "_", character_id or "")
    clean = clean.replace("-", "_").strip("_")
    if not clean:
        clean = "default"
    name = f"cv_{clean}"
    if len(name) < 3:
        name = (name + "___")[:3]
    return name[:512]


def _knowledge_dir(character_dir: str | Path) -> Path:
    return Path(character_dir).expanduser().resolve() / "knowledge"


def _chroma_dir(character_dir: str | Path) -> Path:
    return _knowledge_dir(character_dir) / "chroma"


def _data_relative_path(path: Path) -> Path | None:
    parts = path.parts
    for idx, part in enumerate(parts):
        if part == "data" and idx+1 < len(parts) and parts[idx+1] == "characters":
            return Path(*parts[idx+1:])
    return None


def _settings_from_config(config: dict[str, Any] | None) -> dict[str, Any]:
    pipeline = (config or {}).get("pipeline", {})
    rag = pipeline.get("rag", {}) if isinstance(pipeline, dict) else {}
    return rag if isinstance(rag, dict) else {}


def _embedding_config(config: dict[str, Any] | None) -> tuple[str, dict[str, Any]]:
    inference = (config or {}).get("inference", {})
    inference = inference if isinstance(inference, dict) else {}
    section = inference.get("embedding", {})
    section = section if isinstance(section, dict) else {}
    default = str(section.get("default") or "fake").strip() or "fake"
    provider = section.get(default, {})
    return default, provider if isinstance(provider, dict) else {}


class HashEmbeddings:
    """Small deterministic embedding fallback for tests and offline development."""

    def __init__(self, dimensions: int = 384) -> None:
        self.dimensions = max(16, int(dimensions or 384))

    @staticmethod
    def _tokens(text: str) -> list[str]:
        tokens = re.findall(r"[\w\u4e00-\u9fff]+", text.lower())
        if tokens:
            return tokens
        return [text.lower()] if text else [""]

    def _embed(self, text: str) -> list[float]:
        vec = [0.0] * self.dimensions
        for token in self._tokens(text):
            digest = hashlib.sha256(token.encode("utf-8", errors="ignore")).digest()
            idx = int.from_bytes(digest[:4], "big") % self.dimensions
            sign = -1.0 if digest[4] & 1 else 1.0
            vec[idx] += sign
        norm = math.sqrt(sum(v * v for v in vec)) or 1.0
        return [v / norm for v in vec]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._embed(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed(text)


class RAGEngine:
    """LangChain-backed local RAG index manager.

    Imports LangChain/Chroma lazily so non-RAG inference services can start even
    when optional document dependencies have not been installed yet.
    """

    def __init__(self, config: dict[str, Any] | None = None) -> None:
        self.config = config or {}
        settings = _settings_from_config(self.config)
        self.chunk_size = int(settings.get("chunk_size") or settings.get("chunk_chars") or DEFAULT_CHUNK_SIZE)
        self.chunk_overlap = int(
            settings.get("chunk_overlap") or settings.get("chunk_overlap_chars") or DEFAULT_CHUNK_OVERLAP
        )
        self.default_top_k = int(settings.get("top_k") or DEFAULT_TOP_K)
        self.default_max_context_chars = int(settings.get("max_context_chars") or DEFAULT_MAX_CONTEXT_CHARS)
        self.default_min_score = float(settings.get("min_score") or DEFAULT_MIN_SCORE)
        self._embeddings: Any | None = None

    def _data_root_candidates(self) -> list[Path]:
        candidates: list[Path] = []
        for value in [
            os.getenv("CYBERVERSE_DATA_DIR"),
            str(Path(os.getenv("CYBERVERSE_CONFIG_DIR", "")).expanduser() / "data")
            if os.getenv("CYBERVERSE_CONFIG_DIR")
            else "",
            str(Path.cwd() / "data"),
            str(Path(__file__).resolve().parents[2] / "data"),
        ]:
            value = str(value or "").strip()
            if not value:
                continue
            path = Path(value).expanduser()
            if path not in candidates:
                candidates.append(path)
        return candidates

    def _resolve_source_path(self, source_path: str) -> Path:
        path = Path(source_path).expanduser()
        if path.exists():
            return path.resolve()

        data_relative = _data_relative_path(path)
        if data_relative is not None:
            for data_root in self._data_root_candidates():
                candidate = data_root / data_relative
                if candidate.exists():
                    return candidate.resolve()

        return path.resolve()

    def _documents_cls(self):
        from langchain_core.documents import Document

        return Document

    def _splitter(self):
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        return RecursiveCharacterTextSplitter(
            chunk_size=max(100, self.chunk_size),
            chunk_overlap=max(0, min(self.chunk_overlap, max(0, self.chunk_size - 1))),
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", "；", ";", " ", ""],
        )

    def _embedding_model(self):
        if self._embeddings is not None:
            return self._embeddings

        provider, conf = _embedding_config(self.config)
        if provider == "fake" or conf.get("plugin_class") == "fake":
            self._embeddings = HashEmbeddings(int(conf.get("dimensions") or 384))
            return self._embeddings

        try:
            from langchain_openai import OpenAIEmbeddings
        except Exception as exc:
            raise RuntimeError(
                "RAG embeddings require langchain-openai, or configure inference.embedding.default=fake"
            ) from exc

        api_key = str(conf.get("api_key") or os.getenv("OPENAI_API_KEY") or "")
        base_url = str(conf.get("base_url") or "")
        if provider == "qwen" and not base_url:
            base_url = dashscope_base_url()
            api_key = api_key or os.getenv("DASHSCOPE_API_KEY", "")
        model = str(conf.get("model") or ("text-embedding-v4" if provider == "qwen" else "text-embedding-3-small"))
        kwargs: dict[str, Any] = {"model": model}
        if api_key:
            kwargs["api_key"] = api_key
        if base_url:
            kwargs["base_url"] = base_url
        if conf.get("dimensions"):
            kwargs["dimensions"] = int(conf["dimensions"])
        self._embeddings = OpenAIEmbeddings(**kwargs)
        return self._embeddings

    def _vector_store(self, character_id: str, character_dir: str):
        from langchain_chroma import Chroma

        persist_dir = _chroma_dir(character_dir)
        persist_dir.mkdir(parents=True, exist_ok=True)
        return Chroma(
            collection_name=_safe_collection_name(character_id),
            embedding_function=self._embedding_model(),
            persist_directory=str(persist_dir),
        )

    def _load_text_document(self, path: Path, metadata: dict[str, Any]) -> list[Any]:
        try:
            from langchain_community.document_loaders import TextLoader
        except ImportError as exc:
            raise RuntimeError("RAG text loading requires langchain-community; install cyberverse[rag]") from exc

        return TextLoader(str(path), encoding="utf-8").load()

    def _load_json_documents(self, path: Path, metadata: dict[str, Any]) -> list[Any]:
        Document = self._documents_cls()
        raw = path.read_text(encoding="utf-8")
        try:
            parsed = json.loads(raw)
            text = json.dumps(parsed, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            text = raw
        return [Document(page_content=text, metadata=metadata)]

    def _load_docx_documents(self, path: Path, metadata: dict[str, Any]) -> list[Any]:
        Document = self._documents_cls()
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata=metadata)]

    def _load_documents(self, req: RAGIndexRequest) -> list[Any]:
        path = self._resolve_source_path(req.source_path)
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"source file not found: {path}")

        metadata = {
            "character_id": req.character_id,
            "source_id": req.source_id,
            "source_type": req.source_type,
            "title": req.title,
            "filename": req.filename,
            "mime_type": req.mime_type,
        }
        ext = path.suffix.lower()
        if ext in {".txt", ".md"}:
            docs = self._load_text_document(path, metadata)
        elif ext == ".json":
            docs = self._load_json_documents(path, metadata)
        elif ext == ".pdf":
            try:
                from langchain_community.document_loaders import PyPDFLoader
            except ImportError as exc:
                raise RuntimeError("RAG PDF loading requires langchain-community; install cyberverse[rag]") from exc

            docs = PyPDFLoader(str(path)).load()
        elif ext == ".docx":
            docs = self._load_docx_documents(path, metadata)
        else:
            raise ValueError(f"unsupported knowledge source type: {ext or req.mime_type}")

        for doc in docs:
            doc.metadata = {**metadata, **(doc.metadata or {})}
        return docs

    def _index_source_sync(self, req: RAGIndexRequest) -> int:
        if not req.character_dir:
            raise ValueError("character_dir is required")
        if not req.source_id:
            raise ValueError("source_id is required")

        docs = self._load_documents(req)
        chunks = self._splitter().split_documents(docs)
        chunks = [chunk for chunk in chunks if chunk.page_content.strip()]
        store = self._vector_store(req.character_id, req.character_dir)

        self._delete_source_sync(req.character_id, req.character_dir, req.source_id)
        ids = [f"{req.source_id}:{i}" for i in range(len(chunks))]
        if chunks:
            store.add_documents(chunks, ids=ids)
        return len(chunks)

    def _delete_source_sync(self, character_id: str, character_dir: str, source_id: str) -> None:
        if not character_dir or not source_id:
            return
        persist_dir = _chroma_dir(character_dir)
        if not persist_dir.exists():
            return
        store = self._vector_store(character_id, character_dir)
        collection = getattr(store, "_collection", None)
        if collection is None:
            return
        try:
            collection.delete(where={"source_id": source_id})
        except Exception:
            logger.debug("RAG source delete failed; recreating collection may be required", exc_info=True)

    def _search_sync(self, req: RAGSearchRequest) -> list[RAGSearchResult]:
        query = (req.query or "").strip()
        if not req.character_dir or not query:
            return []
        persist_dir = _chroma_dir(req.character_dir)
        if not persist_dir.exists():
            return []
        store = self._vector_store(req.character_id, req.character_dir)
        top_k = req.top_k if req.top_k > 0 else self.default_top_k
        max_chars = req.max_context_chars if req.max_context_chars > 0 else self.default_max_context_chars
        min_score = req.min_score if req.min_score > 0 else self.default_min_score

        raw_results = store.similarity_search_with_score(query, k=top_k)
        results: list[RAGSearchResult] = []
        used_chars = 0
        for doc, raw_score in raw_results:
            score = 1.0 / (1.0 + max(float(raw_score), 0.0))
            if score < min_score:
                continue
            content = (doc.page_content or "").strip()
            if not content:
                continue
            remaining = max_chars - used_chars
            if remaining <= 0:
                break
            if len(content) > remaining:
                content = content[:remaining]
            used_chars += len(content)
            meta = doc.metadata or {}
            results.append(
                RAGSearchResult(
                    source_id=str(meta.get("source_id") or ""),
                    source_type=str(meta.get("source_type") or ""),
                    title=str(meta.get("title") or ""),
                    filename=str(meta.get("filename") or ""),
                    content=content,
                    score=score,
                )
            )
        return results

    async def index_source(self, req: RAGIndexRequest) -> int:
        return await asyncio.to_thread(self._index_source_sync, req)

    async def delete_source(self, character_id: str, character_dir: str, source_id: str) -> None:
        await asyncio.to_thread(self._delete_source_sync, character_id, character_dir, source_id)

    async def search(self, req: RAGSearchRequest) -> list[RAGSearchResult]:
        return await asyncio.to_thread(self._search_sync, req)

    async def delete_character_index(self, character_dir: str) -> None:
        await asyncio.to_thread(shutil.rmtree, _chroma_dir(character_dir), True)
